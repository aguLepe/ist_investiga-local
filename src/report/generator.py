import streamlit as st
import tempfile
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from graphviz import Source
from sqlalchemy import false


class InformeGenerator:
    def __init__(self):
        # Definición de estilos
        self.estilos = {
            'main': {'font': 'Calibri', 'size': Pt(14), 'color': RGBColor(79, 11, 123), 'bold': True, 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
            'submain': {'font': 'Calibri', 'size': Pt(12), 'color': RGBColor(0, 0, 0), 'bold': False, 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
            'destacado': {'font': 'Calibri', 'size': Pt(12), 'color': RGBColor(0, 0, 0), 'bold': True,
                        'alignment': WD_ALIGN_PARAGRAPH.LEFT},
            'uso': {'font': 'Calibri', 'size': Pt(11), 'color': RGBColor(0, 0, 0), 'bold': False,
                          'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY},
            'encabezado1': {'font': 'Calibri', 'size': Pt(14), 'color': RGBColor(0, 0, 0), 'bold': True, 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
            'encabezado2': {'font': 'Calibri', 'size': Pt(12), 'color': RGBColor(79, 11, 123), 'bold': True, 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
            'encabezado3': {'font': 'Calibri', 'size': Pt(12), 'color': RGBColor(0, 0, 0), 'bold': False,
                            'alignment': WD_ALIGN_PARAGRAPH.LEFT},
            'tabla_header': {'bg_color': '4F0B7B', 'font_color': RGBColor(255, 255, 255), 'bold': True},
        }
        # Datos obligatorios
        self.required_data = {
            'relato': st.session_state.get('relatof'),
            'hechos': st.session_state.get('hechos'),
            'arbol_dot': st.session_state.get('arbol_dot'),
            'medidas': st.session_state.get('edited_measures', [])
        }

    def validate_data(self):
        """Verifica que existan relato, hechos y árbol DOT"""
        faltantes = [k for k, v in self.required_data.items() if not v]

        if faltantes:
            raise ValueError(f"Faltan datos: {', '.join(faltantes)}")
        return True

    def generate_report(self):
        try:
            self.validate_data()
            doc = self._create_document()
            self._add_header(doc)
            self._add_title(doc)
            self._add_portada(doc)
            self._resumen(doc)
            self._metodologia(doc)
            self._add_info_empresa_centro(doc)
            self._add_info_trabajador(doc)
            self._add_detalles_accidente(doc)
            self._fuentes(doc)
            self._add_narrative(doc)
            self._add_facts(doc)
            self._add_cause_tree(doc)
            self._add_corrective_measures(doc)
            self._add_closure(doc)
            return self._save_and_download(doc)
        except Exception as e:
            st.error(f"Error generando informe: {e}")
            return False

    def _create_document(self):
        doc = Document()
        doc.core_properties.language = 'es'
        sec = doc.sections[0]
        sec.top_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)
        return doc

    def _add_header(self, doc):
        sec = doc.sections[0]
        header = sec.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        try:
            run = para.add_run()
            run.add_picture('IST.jpg', width=Cm(2.5))
        except Exception as e:
            st.error(f"Error cargando logo: {e}")

    def _add_title(self, doc):
        titulo = doc.add_heading('Informe Técnico de Investigación', level=1)
        self._apply_style(titulo, 'main')
        cod_p = doc.add_paragraph()
        cod_p.add_run(f"Código: {st.session_state.get('informe_numero','ACC-2024-001')}")
        self._apply_style(cod_p, 'submain')
        doc.add_paragraph()

    def _add_portada(self, doc):
        fa = st.session_state.get('fecha_accidente')
        fecha_acc = fa.strftime('%d/%m/%Y') if hasattr(fa, 'strftime') else str(fa)
        fecha = st.session_state.get('fecha_informe', datetime.today()).strftime('%d/%m/%Y')
        empresa = st.session_state.get('empresa_sel','')
        rut_empresa= st.session_state.get('rut_empresa','')
        self._apply_style(doc.add_paragraph(f'Nombre empresa: {empresa}'),'destacado')
        self._apply_style(doc.add_paragraph(f'RUT empresa: {rut_empresa}'),'destacado')
        self._apply_style(doc.add_paragraph(f'Fecha accidente: {fecha_acc}'),'destacado')
        self._apply_style(doc.add_paragraph(f'Fecha informe: {fecha}'),'destacado')
        doc.add_paragraph()

    def _resumen(self, doc):
        heading = doc.add_heading('1. Resumen del informe', level=2)
        self._apply_style(heading, 'encabezado2')
        doc.add_paragraph()
        self._apply_style(doc.add_paragraph(st.session_state.resumen),'uso')
        doc.add_page_break()

    def _metodologia(self, doc):
        # 1) Encabezado
        heading = doc.add_heading(
            '2. Metodología de análisis de causalidad',
            level=2
        )
        self._apply_style(heading, 'encabezado2')
        # 2) Un pequeño espacio
        doc.add_paragraph()
        # 3) Lista de textos de cada párrafo
        textos = [
            (
                "La metodología utilizada para la presente investigación "
                "se basa en la aplicación de un método denominado 'Árbol de Causas', "
                "el cual, entre otros, es promovido por la OIT y adoptado por el "
                "Ministerio de Salud de Chile para determinar la causalidad de los "
                "accidentes de origen laboral que deriven en consecuencias fatales o graves."
            ),
            (
                "Esta metodología permite, mediante un razonamiento lógico y secuencial, "
                "buscar de manera sistemática los hechos que han estado presentes en la "
                "ocurrencia del accidente y, como tal, facilitar la identificación de "
                "oportunidades de mejoramiento en los procesos de la empresa que se "
                "relacionan principalmente con la situación investigada."
            ),
            (
                "Elemento clave de esta metodología es descartar de forma categórica la "
                "incorporación de juicios de valor como elementos de análisis, considerándose "
                "solo aquellos elementos objetivos identificados y precisados durante el "
                "proceso investigativo —los hechos— que en definitiva corresponden a las "
                "causas que permitieron la ocurrencia del accidente. Estos hechos se "
                "representan gráficamente, lo cual facilita reconocer probables "
                "intervenciones y proponer acciones de mejora."
            ),
            (
                "Es necesario indicar que el modelo utilizado no permite, bajo ningún "
                "concepto o circunstancia, establecer o determinar culpables o responsables "
                "del accidente, sino solo facilitar la identificación de oportunidades de "
                "mejoramiento en los procesos de la empresa."
            )
        ]
        # 4) Crear cada párrafo con estilo 'texto'
        for t in textos:
            self._apply_style(doc.add_paragraph(t), 'uso')
        # 5) Salto de página final
        doc.add_page_break()

    def _add_info_empresa_centro(self, doc):
        heading = doc.add_heading('3. Antecedentes', level=2)
        doc.add_paragraph()
        self._apply_style(heading, 'encabezado2')
        heading = doc.add_heading('3.1. Información de la Empresa y Centro de Trabajo', level=3)
        self._apply_style(heading, 'encabezado3')
        data = [
            ('Razón Social', st.session_state.get('empresa_sel','')),
            ('RUT Empresa', st.session_state.get('rut_empresa','')),
            ('Actividad Económica', st.session_state.get('actividad','')),
            ('Dirección Empresa', st.session_state.get('direccion_empresa','')),
            ('Teléfono Empresa', st.session_state.get('telefono','')),
            ('Representante Legal', st.session_state.get('representante_legal','')),
            ('Región', st.session_state.get('region','')),
            ('Comuna', st.session_state.get('comuna','')),
            ('Centro de Trabajo', st.session_state.get('nombre_local','')),
            ('Dirección Centro', st.session_state.get('direccion_centro',''))
        ]
        self._create_info_table(doc, data)
        doc.add_paragraph()

    def _add_info_trabajador(self, doc):
        heading = doc.add_heading('3.2. Datos del Trabajador', level=3)
        self._apply_style(heading, 'encabezado3')
        fn = st.session_state.get('fecha_nacimiento')
        fecha_nac = fn.strftime('%d/%m/%Y') if hasattr(fn, 'strftime') else str(fn)
        data = [
            ('Nombre Completo', st.session_state.get('nombre_trabajador','')),
            ('RUT Trabajador', st.session_state.get('rut_trabajador','')),
            ('Fecha de Nacimiento', fecha_nac),
            ('Edad', st.session_state.get('edad','')),
            ('Nacionalidad', st.session_state.get('nacionalidad','')),
            ('Estado Civil', st.session_state.get('estado_civil','')),
            ('Tipo de Contrato', st.session_state.get('contrato','')),
            ('Antigüedad en la empresa', st.session_state.get('antiguedad_empresa', '')),
            ('Cargo', st.session_state.get('cargo_trabajador','')),
            ('Antigüedad en el Cargo', st.session_state.get('antiguedad_cargo','')),
            ('Domicilio', st.session_state.get('domicilio',''))
        ]
        self._create_info_table(doc, data)
        doc.add_paragraph()

    def _add_detalles_accidente(self, doc):
        heading = doc.add_heading('3.3. Datos del Accidente', level=3)
        self._apply_style(heading, 'encabezado3')
        fa = st.session_state.get('fecha_accidente')
        fecha_acc = fa.strftime('%d/%m/%Y') if hasattr(fa, 'strftime') else str(fa)
        ha = st.session_state.get('hora_accidente')
        hora_acc = ha.strftime('%H:%M') if hasattr(ha, 'strftime') else str(ha)
        data = [
            ('Fecha', fecha_acc),
            ('Hora', hora_acc),
            ('Lugar', st.session_state.get('lugar_accidente','')),
            ('Tipo', st.session_state.get('tipo_accidente','')),
            ('Naturaleza Lesión', st.session_state.get('naturaleza_lesion','')),
            ('Parte Afectada', st.session_state.get('parte_afectada','')),
            ('Tarea Ejecutada', st.session_state.get('tarea','')),
            ('Operación', st.session_state.get('operacion','')),
            ('Daños a Personas', st.session_state.get('daños_personas','')),
            ('Daños a Propiedad', st.session_state.get('daños_propiedad','')),
            ('Pérdidas en Proceso', st.session_state.get('perdidas_proceso',''))
        ]
        self._create_info_table(doc, data)
        doc.add_page_break()

    def _fuentes(self, doc):
        heading = doc.add_heading('4. Principales fuentes de información', level=2)
        self._apply_style(heading, 'encabezado2')
        doc.add_paragraph()

        # Declaraciones
        heading_fuentes = doc.add_heading('Declaraciones y entrevistas', level=3)
        self._apply_style(heading_fuentes, 'encabezado3')


        if st.session_state.get('declaracion_accidentado'):
            # Datos en línea
            datos_acc = [
                f"Declaración accidentado:",
                f"{st.session_state.get('nombre_trabajador', 'No registrado')},",
                f"RUT {st.session_state.get('rut_trabajador', 'No registrado')},",
                f"{st.session_state.get('cargo_trabajador', 'No registrado')}."
            ]
            doc.add_paragraph(" ".join(datos_acc), style='List Bullet')

            # Texto declaración
            #doc.add_paragraph("Declaraciónes:")
            #doc.add_paragraph(f"» {st.session_state.declaracion_accidentado}", style='List')

        # Testigos
        for i in [1, 2]:
            if st.session_state.get(f'decl{i}_texto'):
                datos_testigo = [
                    f"Declaración testigo:",
                    f"{st.session_state.get(f'decl{i}_nombre', 'No registrado')},",
                    f"RUT {st.session_state.get(f'decl{i}_rut', 'No registrado')},",
                    f"{st.session_state.get(f'decl{i}_cargo', 'No registrado')}."
                ]
                self._apply_style(doc.add_paragraph(" ".join(datos_testigo), style='List Bullet'), 'uso')
                # doc.add_paragraph(f"Declaración: {st.session_state.get(f'decl{i}_texto')}", style='List')

        # Sección de Fuentes
        heading_fuentes = doc.add_heading('Documentos adjuntos y fotografias', level=3)
        self._apply_style(heading_fuentes, 'encabezado3')

        if hasattr(st.session_state, 'file_labels') and st.session_state.file_labels:
            for file_data in st.session_state.file_labels.values():
                tipo = 'Imagen' if file_data['type'].split('/')[-1].upper() in ['PNG', 'JPG', 'JPEG'] \
                    else file_data['type'].split('/')[-1].upper()

                #doc.add_paragraph(
                #    f"{file_data.get('label', 'Sin etiqueta')}, "
                #    f"Nombre archivo: {file_data['file_obj'].name}, "
                #    f"Tipo: {tipo}",
                #    style='List'
                #)

                self._apply_style(doc.add_paragraph(
                    f"{file_data.get('label', 'Sin etiqueta')}",
                    style='List Bullet'
                ), 'uso')


        else:
            doc.add_paragraph("No se adjuntaron documentos de evidencia", style='List')

        doc.add_page_break()

    def _add_narrative(self, doc):
        heading = doc.add_heading('5. Descripción del Accidente', level=2)
        doc.add_paragraph()
        self._apply_style(heading, 'encabezado2')
        for rel in self.required_data['relato'].split('\n'):
            if rel.strip():
                self._apply_style(doc.add_paragraph(rel.strip()), 'uso')
        #p = doc.add_paragraph(self.required_data['relato'])
        #self._apply_style(p, 'uso')
        doc.add_page_break()

    def _add_facts(self, doc):
        heading = doc.add_heading('6. Principales hechos identificados', level=2)
        doc.add_paragraph()
        self._apply_style(heading, 'encabezado2')
        for hecho in self.required_data['hechos'].split('\n'):
            if hecho.strip():
                self._apply_style(doc.add_paragraph(hecho.strip()), 'uso')
        doc.add_page_break()

    def _generate_tree_image(self, output_path):
        img_buf = st.session_state.get('cause_tree_img')
        if img_buf:
            img_buf.seek(0)
            doc.add_heading("Árbol de Causas", level=2)
            # Ajusta el ancho en pulgadas si quieres
            doc.add_picture(img_buf, width=Inches(6))

    def _add_cause_tree(self, doc):
        heading = doc.add_heading('7. Arbol de causas', level=2)
        doc.add_paragraph()
        self._apply_style(heading, 'encabezado2')
        self._apply_style(doc.add_paragraph("Acorde a la información señalada y al análisis de ella, es posible estructurar los hechos identificados según el siguiente árbol de causas:"),'uso')
        doc.add_paragraph()

        if st.session_state.get('cause_tree_img'):
            doc.add_picture(st.session_state['cause_tree_img'], width=Cm(16))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph(style='Caption')
        doc.add_page_break()

    def _add_corrective_measures(self, doc):
        medidas = self.required_data.get('medidas', [])
        if not medidas:
            return
        heading = doc.add_heading('8. Prescripciones', level=2)
        doc.add_paragraph()
        self._apply_style(heading, 'encabezado2')
        # Calcular ancho disponible de la página
        section = doc.sections[0]
        available_width = section.page_width - section.left_margin - section.right_margin
        for m in medidas:
            table = doc.add_table(rows=3, cols=3)
            try:
                table.style = 'Table Grid'
            except Exception:
                pass
            table.allow_autofit = False
            # Fila 1: Tipo (col1), Prioridad (col2-3 merged)
            row1 = table.rows[0].cells
            row1[0].text = f"Tipo: {m.get('tipo','')}"
            merged_prio = row1[1].merge(row1[2])
            merged_prio.text = f"Prioridad: {m.get('prioridad','')}"
            # Fila 2: Plazo (col1-2 merged), Responsable (col3)
            row2 = table.rows[1].cells
            merged_pl = row2[1].merge(row2[2])
            row2[0].text = f"Plazo: {m.get('plazo','')}"
            merged_pl.text = f"Responsable: {m.get('responsable','')}"
            # Fila 2: Descripción (todas 3 fused)
            desc_cells = table.rows[2].cells
            merged_desc = desc_cells[0].merge(desc_cells[1]).merge(desc_cells[2])
            merged_desc.text = f"Descripción: {m.get('descripcion','')}"

            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)

                        # Si la celda está vacía, agregar un run con formato
                        if not paragraph.runs:
                            run = paragraph.add_run()
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
            doc.add_paragraph()

    def _add_closure(self, doc):
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        # 1) Nombre del investigador — negrita y centrado
        investigador = st.session_state.get('investigador', '')
        p_inv = doc.add_paragraph(investigador)
        p_inv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p_inv.runs:
            p_inv.runs[0].bold = True

        # 2) Etiqueta "Responsable Investigación" — centrada
        p_resp = doc.add_paragraph('Consultor IST')
        p_resp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 3) Fecha del informe — alineada a la derecha
        doc.add_paragraph()
        doc.add_paragraph()


    def _create_info_table(self, doc, data):
        table = doc.add_table(rows=len(data), cols=2)
        table.style = 'Table Grid'
        for idx, (label, val) in enumerate(data):
            cells = table.rows[idx].cells
            cells[0].text = f"{label}:"
            cells[1].text = str(val)
            for cell in cells:
                p = cell.paragraphs[0]
                if not p.runs:
                    p.add_run()
                run = p.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
        return table

    def _apply_style(self, element, style_name):
        style = self.estilos.get(style_name, {})
        if hasattr(element, 'runs'):
            if not element.runs:
                element.add_run()
            for run in element.runs:
                self._apply_run_style(run, style)
        if 'alignment' in style:
            element.alignment = style['alignment']

    def _apply_run_style(self, run, style):
        if 'font' in style:
            run.font.name = style['font']
        if 'size' in style:
            run.font.size = style['size']
        if 'color' in style:
            run.font.color.rgb = style['color']
        if 'bold' in style:
            run.font.bold = style['bold']

    def _save_and_download(self, doc):
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, 'informe.docx')
            doc.save(path)
            with open(path, 'rb') as f:
                st.download_button(
                    label='⬇️ Descargar Informe',
                    data=f.read(),
                    file_name='informe_investigacion.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )


def export_to_docx():
    gen = InformeGenerator()
    if gen.generate_report():
        st.success('Informe generado correctamente')
